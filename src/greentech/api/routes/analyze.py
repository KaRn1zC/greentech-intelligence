"""Routes pour l'analyse d'articles par IA.

Endpoint principal POST /analyze qui orchestre la chaine complete :
reception → nettoyage → classification IA → resume SaaS → stockage.

Trois modes d'entree sont supportes :
    - URL (champ `url` du JSON) — le contenu est extrait de la page.
    - Texte brut (champ `texte` du JSON) — le contenu est utilise tel quel.
    - Fichier (upload multipart sur `/analyze/file`) — le texte est extrait
      selon le format (.txt, .md, .pdf, .docx, .html).

Les analyses sont **dispatchees vers une queue Celery + Redis** : l'API
retourne immediatement un ``job_id`` (= Celery task id), le worker dedie
execute le pipeline sequentiellement (1 GPU = 1 inference a la fois).
Le client poll ``GET /analyze/{job_id}`` pour recuperer le statut + resultat.

Avant cette refonte, les jobs etaient stockes en memoire dans un ``dict``
process-local, ce qui posait deux problemes :

1. Tout etait perdu au redemarrage de l'API.
2. Plusieurs requetes simultanees pouvaient saturer le GPU (pas de file
   d'attente, juste un asyncio.create_task() en arriere-plan).

Avec Celery + Redis, les jobs survivent au redemarrage, le worker traite
les analyses sequentiellement (concurrency=1), et l'on peut scaler en
ajoutant des workers sur d'autres GPU.
"""

from __future__ import annotations

import io
import re
import uuid
from pathlib import Path

import httpx
from celery.result import AsyncResult
from fastapi import (
    APIRouter,
    Depends,
    File,
    HTTPException,
    UploadFile,
    status,
)
from loguru import logger

from greentech.api.celery_app import celery_app
from greentech.api.dependencies import get_current_user
from greentech.api.schemas.analysis import (
    AnalysisInput,
    AnalysisJobCreated,
    AnalysisResult,
    AnalysisStatus,
)
from greentech.api.tasks import classify_article as classify_article_task
from greentech.data.storage.models import User

router = APIRouter(prefix="/analyze", tags=["Analyse IA"])

# Contraintes d'upload : on refuse les fichiers > 10 Mo pour proteger le serveur
# (parsing PDF/DOCX en memoire). Les articles depassent rarement quelques dizaines
# de Ko, donc cette limite est largement suffisante.
MAX_UPLOAD_SIZE_BYTES = 10 * 1024 * 1024

# Extensions acceptees avec leur extracteur dedie. On se limite aux formats qui
# contiennent majoritairement du texte exploitable pour une analyse d'article.
ALLOWED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx", ".html", ".htm"}


def _extract_article_from_html(html: str, *, url: str | None = None) -> tuple[str | None, str]:
    """Isole le titre et le corps d'article d'un document HTML via trafilatura.

    Trafilatura combine des heuristiques DOM et ``justext`` pour eliminer le
    boilerplate (menus, footers, cookies, widgets sociaux, scripts, balises
    de navigation). Specialise sur les articles de presse, il evite que le
    modele de classification et le resumeur ne voient que du texte d'interface
    au lieu du contenu reel. En cas d'echec (SPA vide, paywall, page tres
    atypique), on retombe sur une extraction regex minimale pour garantir
    qu'on ait au moins quelque chose a analyser.

    Args:
        html: Document HTML complet, deja decode en texte.
        url: URL d'origine, utilisee par trafilatura pour activer certaines
            heuristiques specifiques a un domaine (optionnel).

    Returns:
        Tuple (titre_ou_None, contenu_texte). Le titre peut etre None si
        aucun ``<title>`` ni metadonnee n'a pu etre detecte.
    """
    from trafilatura import bare_extraction

    titre: str | None = None
    texte: str = ""

    try:
        doc = bare_extraction(
            html,
            url=url,
            include_comments=False,
            include_tables=True,
            favor_precision=True,
            with_metadata=True,
        )
    except Exception as exc:
        logger.warning(f"Extraction trafilatura echouee ({url or 'html brut'}) : {exc}")
        doc = None

    # L'API trafilatura peut retourner un Document dataclass ou un dict selon
    # la version et les options. On gere les deux pour rester robuste a une
    # montee de version mineure.
    if doc is not None:
        if isinstance(doc, dict):
            titre_extrait = doc.get("title") or ""
            texte = (doc.get("text") or doc.get("raw_text") or "").strip()
        else:
            titre_extrait = getattr(doc, "title", "") or ""
            texte = (getattr(doc, "text", None) or getattr(doc, "raw_text", "") or "").strip()
        titre = titre_extrait.strip() or None

    if texte:
        return titre, texte

    # Fallback : trafilatura n'a rien trouve (page atypique, JS-rendered, ...).
    # On garde l'ancienne extraction regex comme filet de securite ; elle est
    # bruyante mais vaut mieux qu'un contenu vide qui ferait echouer le
    # pipeline en aval.
    logger.info(f"Fallback regex pour l'extraction HTML (url={url})")
    title_match = re.search(r"<title[^>]*>(.*?)</title>", html, re.IGNORECASE | re.DOTALL)
    if title_match and not titre:
        titre = title_match.group(1).strip() or None

    text = re.sub(r"<script[^>]*>.*?</script>", "", html, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r"<style[^>]*>.*?</style>", "", text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r"<[^>]+>", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return titre, text


async def _extract_text_from_url(url: str) -> tuple[str, str]:
    """Telecharge une URL et en extrait le titre + le corps d'article.

    Utilise httpx pour la requete HTTP (controle fin du timeout, du
    User-Agent et de la politique de redirection), puis delegue l'extraction
    a ``_extract_article_from_html`` qui s'appuie sur trafilatura pour
    isoler le contenu principal. Le titre extrait est tronque a 500
    caracteres pour tenir dans la contrainte de la colonne ``articles.titre``.

    Args:
        url: URL de l'article a extraire.

    Returns:
        Tuple (titre, contenu_texte) ou le contenu est limite au corps
        d'article. Le titre retombe sur l'URL si aucune metadonnee n'est
        disponible, pour garantir un affichage lisible dans l'UI.

    Raises:
        HTTPException: 422 si l'URL est inaccessible ou si le serveur
            distant a retourne une erreur HTTP.
    """
    try:
        async with httpx.AsyncClient(
            timeout=15.0,
            follow_redirects=True,
            headers={"User-Agent": "GreenTech-Bot/1.0"},
        ) as client:
            response = await client.get(url)
            response.raise_for_status()
    except httpx.HTTPError as e:
        msg = f"Impossible d'acceder a l'URL : {e}"
        raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail=msg)

    html = response.text
    titre, contenu = _extract_article_from_html(html, url=url)
    return (titre or url)[:500], contenu


def _extract_text_from_html_bytes(raw: bytes) -> str:
    """Extrait le corps d'article d'un HTML brut via trafilatura.

    Utilise lors de l'upload d'un fichier ``.html`` / ``.htm`` : la logique
    d'extraction est identique a celle d'une URL, mais sans aller-retour
    reseau. Le titre est ignore ici car l'appelant (``_extract_text_from_upload``)
    se sert du nom de fichier comme titre par defaut.

    Args:
        raw: Contenu brut du fichier HTML.

    Returns:
        Texte du corps d'article debarrasse du boilerplate.
    """
    html = raw.decode("utf-8", errors="replace")
    _, contenu = _extract_article_from_html(html)
    return contenu


def _extract_text_from_pdf(raw: bytes) -> str:
    """Extrait le texte d'un PDF en memoire via pypdf.

    pypdf ne supporte pas les PDF scannes (images) : dans ce cas il retourne
    une chaine vide ou quasi vide. L'appelant doit verifier la longueur.
    """
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(raw))
    parts: list[str] = []
    for page in reader.pages:
        page_text = page.extract_text() or ""
        if page_text.strip():
            parts.append(page_text)
    return "\n\n".join(parts).strip()


def _extract_text_from_docx(raw: bytes) -> str:
    """Extrait le texte d'un DOCX en memoire via python-docx.

    Concatene tous les paragraphes non vides avec un saut de ligne.
    Les tableaux sont traites ligne par ligne, cellule par cellule.
    """
    from docx import Document

    document = Document(io.BytesIO(raw))
    parts: list[str] = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts).strip()


async def _extract_text_from_upload(upload: UploadFile) -> tuple[str, str]:
    """Lit un fichier uploade et extrait titre + contenu textuel.

    Dispatche vers l'extracteur approprie selon l'extension. Le nom de fichier
    (sans extension) sert de titre par defaut, ce qui est plus parlant pour
    l'utilisateur que la troncature des 100 premiers caracteres du contenu.

    Args:
        upload: Fichier fourni par FastAPI via UploadFile.

    Returns:
        Tuple (titre, contenu_texte).

    Raises:
        HTTPException: Si l'extension est refusee, si le fichier depasse la
            taille limite, ou si l'extraction echoue / retourne trop peu de texte.
    """
    filename = upload.filename or "document"
    extension = Path(filename).suffix.lower()

    if extension not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
            detail=(
                f"Extension '{extension}' non supportee. Formats acceptes : "
                f"{', '.join(sorted(ALLOWED_EXTENSIONS))}"
            ),
        )

    raw = await upload.read()
    if len(raw) > MAX_UPLOAD_SIZE_BYTES:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail=f"Fichier trop volumineux (max {MAX_UPLOAD_SIZE_BYTES // (1024 * 1024)} Mo)",
        )
    if not raw:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="Fichier vide",
        )

    try:
        if extension in {".txt", ".md"}:
            contenu = raw.decode("utf-8", errors="replace").strip()
        elif extension == ".pdf":
            contenu = _extract_text_from_pdf(raw)
        elif extension == ".docx":
            contenu = _extract_text_from_docx(raw)
        else:  # .html / .htm
            contenu = _extract_text_from_html_bytes(raw)
    except HTTPException:
        raise
    except Exception as exc:
        logger.exception(f"Extraction echouee pour {filename}")
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"Impossible d'extraire le texte du fichier : {exc}",
        ) from exc

    if len(contenu) < 50:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=(
                "Contenu extrait trop court (minimum 50 caracteres). "
                "Si c'est un PDF scanne, il n'est pas lisible sans OCR."
            ),
        )

    titre = Path(filename).stem or "Document uploade"
    return titre[:500], contenu


def _celery_state_to_status(celery_state: str) -> AnalysisStatus:
    """Convertit un etat brut Celery en ``AnalysisStatus`` metier.

    Mapping :

    * ``PENDING`` : tache enqueue mais pas encore prise par un worker
      (ou job_id inexistant - Celery ne distingue pas les deux cotes
      backend). On verifie l'existence reelle de la tache en amont
      via une autre source quand on a besoin de la distinction.
    * ``STARTED`` : worker a pick la tache, execution en cours
      (necessite ``task_track_started=True`` dans celery_app.py).
    * ``SUCCESS`` : execution terminee avec resultat disponible.
    * ``FAILURE`` : exception levee pendant l'execution.
    * ``RETRY`` / ``REVOKED`` : agrege en ``EN_COURS`` faute de mieux
      (notre pipeline ne fait pas de retry et ne revoke pas).
    """
    mapping = {
        "PENDING": AnalysisStatus.EN_ATTENTE,
        "STARTED": AnalysisStatus.EN_COURS,
        "RETRY": AnalysisStatus.EN_COURS,
        "REVOKED": AnalysisStatus.ERREUR,
        "SUCCESS": AnalysisStatus.TERMINE,
        "FAILURE": AnalysisStatus.ERREUR,
    }
    return mapping.get(celery_state, AnalysisStatus.EN_ATTENTE)


def _build_analysis_result(job_id: uuid.UUID, async_result: AsyncResult) -> AnalysisResult:
    """Construit un ``AnalysisResult`` Pydantic depuis un ``AsyncResult`` Celery.

    Selon l'etat :

    * ``SUCCESS`` : extrait le dict retourne par la tache et le fusionne avec
      le statut TERMINE. La date d'analyse est lue depuis le resultat (la
      tache la set en UTC).
    * ``FAILURE`` : extrait le message d'exception via ``async_result.info``
      (qui contient l'objet Exception levee par la tache).
    * Autres etats : retourne juste statut + job_id, le reste reste None.
    """
    statut = _celery_state_to_status(async_result.state)

    if async_result.state == "SUCCESS":
        payload = async_result.result or {}
        date_str = payload.get("date_analyse")
        date_analyse = None
        if isinstance(date_str, str):
            from datetime import datetime as _dt
            try:
                date_analyse = _dt.fromisoformat(date_str)
            except ValueError:
                date_analyse = None
        return AnalysisResult(
            job_id=job_id,
            statut=statut,
            id_article=payload.get("id_article"),
            titre=payload.get("titre"),
            est_green_it=payload.get("est_green_it"),
            score_confiance=payload.get("score_confiance"),
            resume=payload.get("resume"),
            resume_ecologique=payload.get("resume_ecologique"),
            modele_classification=payload.get("modele_classification"),
            temps_inference_ms=payload.get("temps_inference_ms"),
            date_analyse=date_analyse,
        )

    if async_result.state == "FAILURE":
        return AnalysisResult(
            job_id=job_id,
            statut=statut,
            erreur=str(async_result.info) if async_result.info else "Erreur inconnue",
        )

    return AnalysisResult(job_id=job_id, statut=statut)


@router.post(
    "",
    response_model=AnalysisJobCreated,
    status_code=status.HTTP_202_ACCEPTED,
    summary="Lancer une analyse IA",
)
async def create_analysis(
    data: AnalysisInput,
    _current_user: User = Depends(get_current_user),
) -> AnalysisJobCreated:
    """Soumet une URL ou un texte pour analyse IA via la queue Celery.

    L'analyse est dispatchee vers un worker Celery dedie qui execute le
    pipeline (extraction, classification Qwen3-4B TIES, summarize, stockage).
    L'API retourne immediatement sans bloquer : le client peut soumettre
    plusieurs analyses en serie sans attendre la fin de la precedente.

    Le ``job_id`` retourne est le ``task_id`` Celery (UUID). Il sert a
    consulter le statut/resultat via ``GET /analyze/{job_id}``.

    Args:
        data: URL ou texte a analyser (exactement un des deux).
        _current_user: Utilisateur authentifie (verification du token JWT).

    Returns:
        Identifiant du job + statut initial (en_attente jusqu'a ce qu'un
        worker pick la tache, puis en_cours, puis termine ou erreur).
    """
    async_result = classify_article_task.delay(
        url=data.url,
        texte=data.texte,
        titre_override=None,
    )
    job_id = uuid.UUID(async_result.id)

    source = f"URL={data.url}" if data.url else f"texte ({len(data.texte or '')} chars)"
    logger.info(f"Analyse enqueueed (Celery) {job_id} : {source}")

    return AnalysisJobCreated(job_id=job_id)


@router.post(
    "/file",
    response_model=AnalysisJobCreated,
    status_code=status.HTTP_202_ACCEPTED,
    summary="Lancer une analyse IA sur un fichier uploade",
)
async def create_analysis_from_file(
    fichier: UploadFile = File(..., description="Fichier a analyser"),
    _current_user: User = Depends(get_current_user),
) -> AnalysisJobCreated:
    """Accepte un fichier uploade, extrait son texte et enqueue l'analyse.

    Formats supportes : .txt, .md, .pdf, .docx, .html, .htm.
    La taille est plafonnee a 10 Mo. Les PDF scannes (images) ne sont pas
    lisibles sans OCR et seront rejetes avec une erreur 422.

    L'extraction du texte se fait synchrone dans la requete HTTP (rapide,
    sans GPU), puis le texte est envoye dans la queue Celery comme pour
    le mode texte brut. Le client recoit immediatement un ``job_id``.

    Args:
        fichier: Fichier uploade (multipart/form-data).
        _current_user: Utilisateur authentifie (verification du token JWT).

    Returns:
        Identifiant du job d'analyse + statut initial.

    Raises:
        HTTPException: 413 si fichier trop gros, 415 si format non supporte,
            422 si le contenu extrait est insuffisant ou le fichier illisible.
    """
    titre, contenu = await _extract_text_from_upload(fichier)

    async_result = classify_article_task.delay(
        url=None,
        texte=contenu,
        titre_override=titre,
    )
    job_id = uuid.UUID(async_result.id)

    logger.info(
        f"Analyse fichier enqueueed (Celery) {job_id} : "
        f"nom='{fichier.filename}' taille={len(contenu)} chars"
    )
    return AnalysisJobCreated(job_id=job_id)


@router.get(
    "/{job_id}",
    response_model=AnalysisResult,
    summary="Statut d'une analyse",
)
async def get_analysis_status(
    job_id: uuid.UUID,
    _current_user: User = Depends(get_current_user),
) -> AnalysisResult:
    """Retourne le statut et le resultat d'un job d'analyse.

    Lit l'etat de la tache Celery via le result backend Redis :

    * ``en_attente`` : tache dans la queue, pas encore prise par un worker.
    * ``en_cours`` : worker en train d'executer le pipeline.
    * ``termine`` : resultat disponible (id_article, est_green_it, resume...).
    * ``erreur`` : exception levee pendant l'execution (champ erreur peuple).

    Args:
        job_id: Identifiant UUID du job retourne par POST /analyze.
        _current_user: Utilisateur authentifie.

    Returns:
        Statut courant et resultats (si termines).
    """
    async_result = AsyncResult(str(job_id), app=celery_app)
    return _build_analysis_result(job_id, async_result)
